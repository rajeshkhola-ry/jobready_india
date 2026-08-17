#!/usr/bin/env python3
"""Builds a Word (.docx) document with one full-page image per page.

Usage: rasterize_to_docx.py <output.docx> <image1> [<image2> ...]

Final "Tier 3" fallback for scanned/image-only PDFs where neither
pdf2docx nor LibreOffice's native PDF import could produce a usable
DOCX. Uses python-docx directly instead of LibreOffice's HTML import,
which can fail to embed file:// or relative image URIs in headless
mode - this guarantees the images are actually embedded.
Page size is derived from the first rasterized image's own pixel
dimensions (at the known rasterization DPI) so each page image fits
its own page exactly, instead of spilling onto extra blank pages or
leaving unused whitespace under a fixed Letter-sized assumption.
Prints a JSON result to stdout: {"success": true, "pages": N} or
{"error": "<message>"}.
"""

import sys
import json

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_BREAK
except ImportError as exc:
    print(json.dumps({'error': 'python-docx not installed: ' + str(exc)}))
    sys.exit(1)

try:
    from PIL import Image
except ImportError:
    Image = None

RASTER_DPI = 150
MARGIN_INCHES = 0.25
FALLBACK_CONTENT_WIDTH_INCHES = 6.5
MAX_PAGE_WIDTH_INCHES = 11.0
MAX_PAGE_HEIGHT_INCHES = 17.0


def _resolve_page_geometry(first_image_path):
    """Returns (content_width_in, page_width_in, page_height_in) or (None, None, None)."""
    if Image is None:
        return None, None, None

    try:
        with Image.open(first_image_path) as source:
            px_width, px_height = source.size
    except Exception:
        return None, None, None

    page_width_in = min(MAX_PAGE_WIDTH_INCHES, px_width / RASTER_DPI + MARGIN_INCHES * 2)
    page_height_in = min(MAX_PAGE_HEIGHT_INCHES, px_height / RASTER_DPI + MARGIN_INCHES * 2)
    return page_width_in - MARGIN_INCHES * 2, page_width_in, page_height_in


def main():
    if len(sys.argv) < 3:
        print(json.dumps({'error': 'Usage: rasterize_to_docx.py <output.docx> <image1> [<image2> ...]'}))
        sys.exit(1)

    output_path = sys.argv[1]
    image_paths = sys.argv[2:]

    try:
        document = Document()
        content_width_in, page_width_in, page_height_in = _resolve_page_geometry(image_paths[0])

        if page_width_in and page_height_in:
            section = document.sections[0]
            section.page_width = Inches(page_width_in)
            section.page_height = Inches(page_height_in)
            section.left_margin = Inches(MARGIN_INCHES)
            section.right_margin = Inches(MARGIN_INCHES)
            section.top_margin = Inches(MARGIN_INCHES)
            section.bottom_margin = Inches(MARGIN_INCHES)
        else:
            content_width_in = FALLBACK_CONTENT_WIDTH_INCHES

        for index, image_path in enumerate(image_paths):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Inches(0)
            paragraph.paragraph_format.space_after = Inches(0)
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(content_width_in))
            if index < len(image_paths) - 1:
                paragraph.add_run().add_break(WD_BREAK.PAGE)

        document.save(output_path)
    except Exception as exc:
        print(json.dumps({'error': 'Failed to build DOCX: ' + str(exc)}))
        sys.exit(1)

    print(json.dumps({'success': True, 'pages': len(image_paths)}))


if __name__ == '__main__':
    main()
